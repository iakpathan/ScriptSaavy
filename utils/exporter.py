from reportlab.lib.pagesizes import letter
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
from reportlab.lib.styles import getSampleStyleSheet
from docx import Document
import os
import subprocess


from fpdf import FPDF


from datetime import datetime
from fpdf import FPDF


def export_script(content, format_type='pdf'):
    """Export script to the desired format (PDF or text)."""
    timestamp = datetime.now().strftime("%Y%m%d%H%M%S")
    export_dir = "exports"
    os.makedirs(export_dir, exist_ok=True)
    file_path = f"{export_dir}/script_{timestamp}.{format_type}"

    if format_type == 'pdf':
        # ✅ Export as PDF with UTF-8 support
        doc = SimpleDocTemplate(file_path, pagesize=letter)
        styles = getSampleStyleSheet()
        content_utf8 = content.encode("utf-8", "ignore").decode("utf-8")
        story = [Paragraph(content_utf8, styles["BodyText"])]
        doc.build(story)
    else:
        # ✅ Export as text with UTF-8 encoding
        with open(file_path, "w", encoding="utf-8") as f:
            f.write(content)

    return file_path






def open_pdf(export_path):
    """Open the generated PDF file using the default system viewer."""
    if os.path.exists(export_path):
        print(f"✅ File '{export_path}' generated successfully!")

        try:
            # Open based on OS
            if os.name == 'nt':  # Windows
                os.startfile(export_path)
            elif os.name == 'posix':
                if 'darwin' in os.sys.platform:  # macOS
                    subprocess.run(["open", export_path], check=True)
                else:  # Linux
                    if subprocess.run(["which", "xdg-open"], capture_output=True).returncode == 0:
                        subprocess.run(["xdg-open", export_path], check=True)
                    else:
                        print("⚠️ 'xdg-open' not found. Please open the file manually.")
        except Exception as e:
            print(f"❌ Error opening file: {str(e)}")
    else:
        print(f"❌ File '{export_path}' not found. Check your export function.")


def export_to_pdf(script):
    """
    Export the script to a PDF file.

    Args:
        script (str): Script content to be saved in PDF format.

    Returns:
        str: Path to the generated PDF file.
    """
    export_path = "generated_script.pdf"

    try:
        doc = SimpleDocTemplate(export_path, pagesize=letter)
        styles = getSampleStyleSheet()
        content = []

        # Split and process the script correctly
        for line in script.split('\n'):
            if line.strip():  # Skip empty lines
                content.append(Paragraph(line, styles['BodyText']))
                content.append(Spacer(1, 12))

        # Build the PDF only if there’s content
        if content:
            doc.build(content)
        else:
            raise ValueError("❌ Script is empty. Cannot generate PDF.")

    except Exception as e:
        raise RuntimeError(f"❌ Failed to generate PDF: {str(e)}")

    return export_path


def export_to_docx(script):
    """
    Export the script to a DOCX file.

    Args:
        script (str): Script content to be saved in DOCX format.

    Returns:
        str: Path to the generated DOCX file.
    """
    export_path = "generated_script.docx"

    try:
        doc = Document()

        # Split and process the script correctly
        for line in script.split('\n'):
            if line.strip():  # Skip empty lines
                doc.add_paragraph(line)

        # Save the DOCX file only if content exists
        if len(doc.paragraphs) > 0:
            doc.save(export_path)
        else:
            raise ValueError("❌ Script is empty. Cannot generate DOCX.")

    except Exception as e:
        raise RuntimeError(f"❌ Failed to generate DOCX: {str(e)}")

    return export_path


def delete_previous_files():
    """
    Deletes previously generated script files (PDF and DOCX)
    to ensure new content replaces them.
    """
    files_to_delete = ["generated_script.pdf", "generated_script.docx"]

    for file in files_to_delete:
        if os.path.exists(file):
            os.remove(file)
            print(f"🗑️ Deleted old file: {file}")
        else:
            print(f"⚠️ No previous file found: {file}")


# Example at the end of exporter.py (for testing purposes)
if __name__ == "__main__":
    script_content = "Here is a sample script generated by the AI.\nThis is a test for exporting to PDF and DOCX."

    # Export the script to PDF and open it
    export_path = export_script(script_content, format_type='pdf')
    open_pdf(export_path)
